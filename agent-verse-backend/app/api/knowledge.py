"""Knowledge API — collections, document ingestion, hybrid search, semantic cache."""

from __future__ import annotations

import hashlib
import uuid as _uuid
from typing import Any

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile, status
from pydantic import BaseModel, SecretStr

from app.rag.models import Chunk, Document, KnowledgeCollection
from app.rag.semantic_cache import SemanticCache
from app.rag.store import KnowledgeStore
from app.tenancy.context import TenantContext

router = APIRouter(prefix="/knowledge", tags=["knowledge"])

# Embedding dimension used for random dummy embeddings when no real embedder is present.
_EMBEDDING_DIM = 768


# ---------------------------------------------------------------------------
# Request models
# ---------------------------------------------------------------------------

class CreateCollectionRequest(BaseModel):
    name: str
    description: str = ""
    embedder_type: str = "voyage"


class IngestRequest(BaseModel):
    collection_id: str
    source_type: str = "text"  # git / markdown / text / openapi / python / code
    content: str
    metadata: dict[str, Any] = {}


class RepoIngestRequest(BaseModel):
    repo_url: str
    collection_id: str
    branch: str = "main"
    file_patterns: list[str] = ["**/*.py", "**/*.md", "**/*.ts", "**/*.js"]
    max_files: int = 200


class OpenAPIIngestRequest(BaseModel):
    content: str  # OpenAPI JSON or YAML string
    collection_id: str
    source_url: str = ""


class UrlIngestRequest(BaseModel):
    collection_id: str
    url: str
    source_type: str = "web"  # web|github|confluence|jira|slack


class GitHubIngestRequest(BaseModel):
    collection_id: str
    owner: str
    repo: str
    branch: str = "HEAD"
    max_files: int = 300


class ConfluenceIngestRequest(BaseModel):
    collection_id: str
    base_url: str
    space_key: str
    token: SecretStr  # SecretStr prevents token from appearing in logs or tracebacks
    user: str
    max_pages: int = 1000


class JiraIngestRequest(BaseModel):
    collection_id: str
    base_url: str
    project_key: str
    token: SecretStr  # SecretStr prevents token from appearing in logs or tracebacks
    user: str
    jql_extra: str = ""
    max_issues: int = 500


class SlackIngestRequest(BaseModel):
    collection_id: str
    channel_id: str
    token: SecretStr  # SecretStr prevents token from appearing in logs or tracebacks
    channel_name: str = ""
    max_messages: int = 500


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _require_tenant(request: Request) -> Any:
    ctx = getattr(request.state, "tenant", None)
    if ctx is None:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Unauthorized")
    return ctx


def _knowledge_store(request: Request) -> KnowledgeStore:
    return request.app.state.knowledge_store  # type: ignore[no-any-return]


def _semantic_cache(request: Request) -> SemanticCache:
    return request.app.state.semantic_cache  # type: ignore[no-any-return]


def _cache_stats(request: Request) -> dict[str, dict[str, int]]:
    """Per-tenant hit/miss counters stored lazily on app.state."""
    if not hasattr(request.app.state, "_cache_stats"):
        request.app.state._cache_stats = {}
    return request.app.state._cache_stats  # type: ignore[no-any-return]


def _fallback_embedding(dim: int = _EMBEDDING_DIM) -> list[float]:
    raise HTTPException(
        status_code=503,
        detail=(
            "Embedding provider not configured. "
            "Set one of: VOYAGE_API_KEY, OPENAI_API_KEY, GOOGLE_API_KEY "
            "(cloud embeddings) or SENTENCE_TRANSFORMERS_MODEL=all-MiniLM-L6-v2 "
            "(local CPU embeddings via sentence-transformers)."
        )
    )


# ---------------------------------------------------------------------------
# Endpoints — collections
# ---------------------------------------------------------------------------

@router.get("/collections")
async def list_collections(request: Request) -> list[dict[str, Any]]:
    tenant_ctx: TenantContext = _require_tenant(request)
    store = _knowledge_store(request)
    collections = store.list_collections(tenant_ctx=tenant_ctx)
    return [
        {
            "collection_id": c.collection_id,
            "name": c.name,
            "description": c.description,
            "document_count": c.document_count,
            "embedder": c.embedder,
        }
        for c in collections
    ]


@router.post("/collections", status_code=status.HTTP_201_CREATED)
async def create_collection(
    request: Request, body: CreateCollectionRequest
) -> dict[str, Any]:
    tenant_ctx: TenantContext = _require_tenant(request)
    store = _knowledge_store(request)
    collection = KnowledgeCollection(
        name=body.name,
        description=body.description,
        embedder=body.embedder_type,
    )
    cid = store.create_collection(collection, tenant_ctx=tenant_ctx)
    return {
        "collection_id": cid,
        "name": body.name,
        "description": body.description,
        "document_count": 0,
        "embedder": body.embedder_type,
    }


@router.delete("/collections/{collection_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_collection(request: Request, collection_id: str) -> None:
    tenant_ctx: TenantContext = _require_tenant(request)
    store = _knowledge_store(request)

    # Verify collection exists and belongs to this tenant
    collection = store.get_collection(collection_id, tenant_ctx=tenant_ctx)
    if collection is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Collection {collection_id} not found",
        )

    # H-5: Block deletion if a legal hold is active on this resource
    from app.governance.legal_holds import LegalHoldManager
    _legal_hold_mgr = getattr(request.app.state, "legal_hold_manager", None)
    if _legal_hold_mgr is not None:
        try:
            _is_held = await _legal_hold_mgr.is_under_hold(
                resource_id=collection_id, tenant_id=tenant_ctx.tenant_id
            )
            if _is_held:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="Resource is under legal hold and cannot be deleted",
                )
        except HTTPException:
            raise
        except Exception:
            pass  # Legal hold check failure is non-fatal; allow deletion

    # Delete from DB: chunks first (FK constraint), then the collection row
    db = getattr(store, "_db", None)
    if db is not None:
        try:
            from sqlalchemy import text
            async with db() as session, session.begin():
                await session.execute(
                    text(
                        "DELETE FROM documents "
                        "WHERE collection_id = :cid AND tenant_id = :tid"
                    ),
                    {"cid": collection_id, "tid": tenant_ctx.tenant_id},
                )
                await session.execute(
                    text(
                        "DELETE FROM knowledge_collections "
                        "WHERE id = :cid AND tenant_id = :tid"
                    ),
                    {"cid": collection_id, "tid": tenant_ctx.tenant_id},
                )
        except Exception as exc:
            import logging
            logging.getLogger(__name__).warning("delete_collection_db_failed: %s", exc)

    # Remove from in-memory cache
    key = (tenant_ctx.tenant_id, collection_id)
    store._data.pop(key, None)  # type: ignore[attr-defined]


# ---------------------------------------------------------------------------
# Endpoints — ingestion
# ---------------------------------------------------------------------------

@router.post("/ingest", status_code=status.HTTP_201_CREATED)
async def ingest_document(
    request: Request, body: IngestRequest
) -> dict[str, Any]:
    tenant_ctx: TenantContext = _require_tenant(request)
    store = _knowledge_store(request)

    # Verify collection exists.
    collection = store.get_collection(body.collection_id, tenant_ctx=tenant_ctx)
    if collection is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Collection {body.collection_id} not found",
        )

    content_hash = hashlib.sha256(body.content.encode()).hexdigest()
    str_metadata = {k: str(v) for k, v in body.metadata.items()}
    document = Document(
        collection_id=body.collection_id,
        source=body.source_type,
        content=body.content,
        content_hash=content_hash,
        metadata=str_metadata,
    )

    # Split into token-aware chunks for accurate LLM context window usage.
    from app.knowledge.chunker_v2 import chunk_by_tokens
    _raw_chunks = chunk_by_tokens(body.content, max_tokens=512, overlap_tokens=64)
    chunks_text = [type("_C", (), {"content": c, "start_char": 0, "end_char": len(c)})() for c in _raw_chunks]

    # Fallback: very short content that doesn't meet min_chunk threshold
    if not chunks_text and body.content.strip():
        chunks_text = [type("_C", (), {"content": body.content.strip(), "start_char": 0, "end_char": len(body.content)})()]

    chunks_created = 0
    embedder = getattr(request.app.state, "embedder", None)
    from app.providers.base import embed_texts
    for idx, text_chunk in enumerate(chunks_text):
        chunk_text = text_chunk.content
        embeddings = await embed_texts([chunk_text], provider=embedder)
        chunk_embedding = embeddings[0]
        chunk = Chunk(
            document_id=document.document_id,
            content=chunk_text,
            embedding=chunk_embedding,
            chunk_index=idx,
            metadata={**str_metadata, "source_type": body.source_type},
        )
        store.ingest_chunk(chunk, collection_id=body.collection_id, tenant_ctx=tenant_ctx)
        chunks_created += 1

    return {
        "document_id": document.document_id,
        "collection_id": body.collection_id,
        "chunks_created": chunks_created,
        "content_hash": content_hash,
    }


# ---------------------------------------------------------------------------
# Endpoints — search
# ---------------------------------------------------------------------------

@router.get("/search")
async def search_knowledge(
    request: Request,
    q: str,
    collection_id: str,
    top_k: int = 10,
    threshold: float = 0.5,
) -> list[dict[str, Any]]:
    tenant_ctx: TenantContext = _require_tenant(request)
    store = _knowledge_store(request)

    embedder = getattr(request.app.state, "embedder", None)

    # FIX 5: Fail loudly when no embedder is configured.
    # Previously: returned empty embeddings silently, corrupting search results.
    # Now: raises 503 with an actionable message for the operator.
    if embedder is None:
        raise HTTPException(
            status_code=503,
            detail=(
                "No embedding provider configured. "
                "Set VOYAGE_API_KEY or OPENAI_API_KEY to enable knowledge base search."
            ),
        )

    from app.providers.base import embed_texts
    query_embeddings = await embed_texts([q], provider=embedder)
    query_embedding = query_embeddings[0]
    if hasattr(store, "hybrid_search_db"):
        results = await store.hybrid_search_db(
            q, query_embedding, collection_id, tenant_ctx, top_k=top_k
        )
    else:
        results = store.hybrid_search(q, query_embedding, collection_id, tenant_ctx, top_k=top_k)
    return [
        {
            "chunk_id": r.chunk_id,
            "content": r.content,
            "score": r.score,
            "vector_score": r.vector_score,
            "trigram_score": r.trigram_score,
            # Source citation fields
            "source_file": getattr(r, "metadata", {}).get("source_file", ""),
            "source_url": getattr(r, "metadata", {}).get("source_url", ""),
            "char_offset": getattr(r, "metadata", {}).get("char_offset"),
            "line_start": getattr(r, "metadata", {}).get("line_start"),
        }
        for r in results
        if r.score >= threshold
    ]


# ---------------------------------------------------------------------------
# Endpoints — semantic cache
# ---------------------------------------------------------------------------

@router.get("/cache/stats")
async def get_cache_stats(request: Request) -> dict[str, Any]:
    tenant_ctx: TenantContext = _require_tenant(request)
    cache = _semantic_cache(request)
    stats = cache.stats(tenant_ctx=tenant_ctx)
    return {
        "tenant_id": tenant_ctx.tenant_id,
        "hits": stats["hits"],
        "misses": stats["misses"],
    }


@router.delete("/cache", status_code=status.HTTP_204_NO_CONTENT)
async def clear_cache(request: Request) -> None:
    tenant_ctx: TenantContext = _require_tenant(request)
    cache = _semantic_cache(request)
    cache.clear(tenant_ctx=tenant_ctx)


# ---------------------------------------------------------------------------
# Endpoints — file upload ingestion
# ---------------------------------------------------------------------------

@router.post("/ingest/file", status_code=201)
async def ingest_file(
    request: Request,
    file: UploadFile = File(...),
    collection_id: str = Form(...),
) -> dict[str, Any]:
    """Ingest a file into a knowledge collection.

    Supports: .txt, .md, .py, .ts, .js, .json, .pdf, .docx
    Open source parsing only (pypdf, python-docx if installed).
    """
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    content_bytes = await file.read()
    filename = file.filename or "uploaded_file"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"
    source_type = "code" if ext in {"py", "ts", "js", "jsx", "tsx"} else "text"

    # Parse content
    if ext == "pdf":
        try:
            import io

            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            text = content_bytes.decode("utf-8", errors="replace")
    elif ext in {"docx", "doc"}:
        try:
            import io

            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            text = "\n".join(para.text for para in doc.paragraphs)
        except ImportError:
            text = content_bytes.decode("utf-8", errors="replace")
    else:
        text = content_bytes.decode("utf-8", errors="replace")

    if not text.strip():
        raise HTTPException(422, "File is empty or could not be parsed")

    # Chunk using token-aware chunker
    from app.knowledge.chunker_v2 import chunk_by_tokens as _chunk_by_tokens_file
    _raw_file_chunks = _chunk_by_tokens_file(text, max_tokens=512, overlap_tokens=64)
    chunks = [type("_C", (), {"content": c, "start_char": 0, "end_char": len(c)})() for c in _raw_file_chunks]

    # Fallback for very short content
    if not chunks and text.strip():
        chunks = [type("_C", (), {"content": text.strip(), "start_char": 0, "end_char": len(text)})()]

    ingested = 0
    document_id = _uuid.uuid4().hex
    from app.providers.base import EmbedRequest
    for idx, chunk in enumerate(chunks):
        if not chunk.content.strip():
            continue
        embedding: list[float] = []
        if embedder:
            try:
                resp = await embedder.embed(EmbedRequest(texts=[chunk.content]))
                embedding = resp.embeddings[0] if resp.embeddings else []
            except Exception:
                embedding = []
        rag_chunk = Chunk(
            document_id=document_id,
            content=chunk.content,
            embedding=embedding,
            chunk_index=idx,
            metadata={
                "source_file": filename,
                "ext": ext,
                "char_offset": str(chunk.start_char),
                "source_type": source_type,
            },
        )
        store.ingest_chunk(rag_chunk, collection_id=collection_id, tenant_ctx=tenant)
        ingested += 1

    return {
        "filename": filename,
        "chunks_created": ingested,
        "collection_id": collection_id,
        "file_size_bytes": len(content_bytes),
    }


# ---------------------------------------------------------------------------
# Endpoints — repository ingestion
# ---------------------------------------------------------------------------

@router.post("/ingest/repo", status_code=202)
async def ingest_repository(
    request: Request, body: RepoIngestRequest
) -> dict[str, Any]:
    """Clone a git repository and ingest all matching files.

    Uses git (open source) for cloning. No cloud API calls.
    Returns immediately — ingestion runs in background.
    """
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    import asyncio
    # Run in background task
    task = asyncio.create_task(
        _ingest_repo_background(
            repo_url=body.repo_url,
            collection_id=body.collection_id,
            branch=body.branch,
            file_patterns=body.file_patterns,
            max_files=body.max_files,
            store=store,
            embedder=embedder,
            tenant_ctx=tenant,
        )
    )
    # Don't await — return immediately
    _ = task  # Task runs in background

    return {
        "status": "ingestion_started",
        "repo_url": body.repo_url,
        "collection_id": body.collection_id,
        "branch": body.branch,
        "message": "Repository ingestion started in background. "
                   "Check /knowledge/collections for progress.",
    }


async def _ingest_repo_background(
    repo_url: str,
    collection_id: str,
    branch: str,
    file_patterns: list[str],
    max_files: int,
    store: Any,
    embedder: Any,
    tenant_ctx: Any,
) -> None:
    import asyncio
    import pathlib
    import shutil
    import tempfile

    from app.observability.logging import get_logger
    from app.providers.base import EmbedRequest
    from app.knowledge.chunker_v2 import chunk_by_tokens as _chunk_by_tokens_repo
    logger = get_logger(__name__)

    tmpdir = tempfile.mkdtemp(prefix="agentverse_repo_")
    try:
        # Clone using git — non-blocking async subprocess
        proc = await asyncio.create_subprocess_exec(
            "git", "clone", "--depth=1", "--branch", branch, repo_url, tmpdir,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        try:
            _stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)
        except TimeoutError:
            proc.kill()
            logger.warning("repo_clone_timeout", repo=repo_url)
            return

        if proc.returncode != 0:
            logger.warning(
                "repo_clone_failed",
                repo=repo_url,
                error=(stderr or b"").decode("utf-8", errors="replace")[:200],
            )
            return

        chunker = None  # replaced by chunk_by_tokens below
        files_processed = 0

        for pattern in file_patterns:
            for filepath in pathlib.Path(tmpdir).rglob(pattern.lstrip("*/")):
                if files_processed >= max_files:
                    break
                if not filepath.is_file():
                    continue
                try:
                    text = filepath.read_text(encoding="utf-8", errors="replace")
                    if not text.strip():
                        continue
                    ext = filepath.suffix.lstrip(".")
                    src_type = "code" if ext in {"py", "ts", "js", "jsx", "tsx"} else "text"
                    _raw = _chunk_by_tokens_repo(text, max_tokens=512, overlap_tokens=64)
                    chunks = [type("_C", (), {"content": c, "start_char": 0, "end_char": len(c)})() for c in _raw]
                    rel_path = str(filepath.relative_to(tmpdir))
                    doc_id = _uuid.uuid4().hex
                    for idx, chunk in enumerate(chunks):
                        embedding: list[float] = []
                        if embedder:
                            try:
                                resp = await embedder.embed(EmbedRequest(texts=[chunk.content]))
                                embedding = resp.embeddings[0] if resp.embeddings else []
                            except Exception:
                                pass
                        rag_chunk = Chunk(
                            document_id=doc_id,
                            content=chunk.content,
                            embedding=embedding,
                            chunk_index=idx,
                            metadata={
                                "source_file": rel_path,
                                "repo_url": repo_url,
                                "char_offset": str(chunk.start_char),
                                "source_type": src_type,
                            },
                        )
                        store.ingest_chunk(rag_chunk, collection_id=collection_id,
                                           tenant_ctx=tenant_ctx)
                    files_processed += 1
                except Exception as exc:
                    logger.warning("repo_file_ingest_failed",
                                   file=str(filepath), error=str(exc))

        logger.info("repo_ingest_complete", repo=repo_url, files=files_processed)
    except Exception as exc:
        logger.warning("repo_ingest_failed", repo=repo_url, error=str(exc))
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ---------------------------------------------------------------------------
# Endpoints — OpenAPI spec ingestion
# ---------------------------------------------------------------------------

@router.post("/ingest/openapi", status_code=201)
async def ingest_openapi(
    request: Request, body: OpenAPIIngestRequest
) -> dict[str, Any]:
    """Ingest an OpenAPI spec — creates a chunk per endpoint."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    try:
        import json as _json
        try:
            spec = _json.loads(body.content)
        except _json.JSONDecodeError:
            import yaml as _yaml
            spec = _yaml.safe_load(body.content)
    except Exception as exc:
        raise HTTPException(422, f"Could not parse OpenAPI spec: {exc}")

    paths = spec.get("paths", {})
    chunks_created = 0
    from app.providers.base import EmbedRequest

    for path, methods in paths.items():
        if not isinstance(methods, dict):
            continue
        for method, op in methods.items():
            if method.lower() not in {"get", "post", "put", "patch", "delete", "options"}:
                continue
            if not isinstance(op, dict):
                continue
            summary = op.get("summary", "")
            description = op.get("description", "")
            params = [p.get("name", "") for p in op.get("parameters", [])
                      if isinstance(p, dict)]

            chunk_text = (
                f"{method.upper()} {path}\n"
                f"Summary: {summary}\n"
                f"Description: {description}\n"
                f"Parameters: {', '.join(params) if params else 'none'}"
            ).strip()

            if not chunk_text:
                continue

            embedding: list[float] = []
            if embedder:
                try:
                    resp = await embedder.embed(EmbedRequest(texts=[chunk_text]))
                    embedding = resp.embeddings[0] if resp.embeddings else []
                except Exception:
                    pass

            rag_chunk = Chunk(
                document_id=_uuid.uuid4().hex,
                content=chunk_text,
                embedding=embedding,
                chunk_index=chunks_created,
                metadata={
                    "source_url": body.source_url,
                    "source_type": "openapi",
                    "endpoint": f"{method.upper()} {path}",
                },
            )
            store.ingest_chunk(rag_chunk, collection_id=body.collection_id,
                               tenant_ctx=tenant)
            chunks_created += 1

    return {
        "endpoints_ingested": chunks_created,
        "collection_id": body.collection_id,
        "source_url": body.source_url,
    }


# ---------------------------------------------------------------------------
# Endpoints — URL-based connector ingestion (Phase 9)
# ---------------------------------------------------------------------------

@router.post("/ingest/url", status_code=201)
async def ingest_from_url(request: Request, body: UrlIngestRequest) -> dict[str, Any]:
    """Ingest content from a URL (web page, GitHub file, Confluence page, etc.)."""
    tenant_ctx = _require_tenant(request)
    store = _knowledge_store(request)

    content = ""
    metadata: dict[str, Any] = {"source_url": body.url, "source_type": body.source_type}

    try:
        if body.source_type == "web":
            import httpx
            async with httpx.AsyncClient(timeout=30.0) as client:
                resp = await client.get(body.url, headers={"User-Agent": "AgentVerse/1.0"})
                resp.raise_for_status()
                raw = resp.text
                import re
                content = re.sub(r'<[^>]+>', ' ', raw)
                content = re.sub(r'\s+', ' ', content).strip()[:50000]
                title_match = re.search(r'<title[^>]*>(.*?)</title>', raw, re.IGNORECASE)
                metadata["title"] = title_match.group(1) if title_match else body.url

        elif body.source_type == "github":
            raw_url = body.url.replace("github.com", "raw.githubusercontent.com").replace("/blob/", "/")
            import httpx
            headers: dict[str, str] = {}
            import os as _os
            if (token := _os.getenv("GITHUB_TOKEN")):
                headers["Authorization"] = f"Bearer {token}"
            async with httpx.AsyncClient(timeout=15.0) as client:
                resp = await client.get(raw_url, headers=headers)
                resp.raise_for_status()
                content = resp.text[:100000]
            metadata["filename"] = body.url.split("/")[-1]

        else:
            raise HTTPException(400, f"Source type '{body.source_type}' not yet supported for URL ingestion. Supported: web, github")

    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, f"Failed to fetch content from URL: {exc}")

    if not content.strip():
        raise HTTPException(422, "No content extracted from URL")

    # Chunk and ingest
    from app.knowledge.chunker_v2 import chunk_by_tokens as _chunk_by_tokens_url
    _raw_url_chunks = _chunk_by_tokens_url(content, max_tokens=512, overlap_tokens=64)
    chunks = [type("_C", (), {"content": c, "start_char": 0, "end_char": len(c)})() for c in _raw_url_chunks]
    if not chunks and content.strip():
        chunks = [type("_C", (), {"content": content.strip(), "start_char": 0, "end_char": len(content)})()]

    embedder = getattr(request.app.state, "embedder", None)
    import uuid as _uuid_mod

    from app.providers.base import embed_texts
    from app.rag.models import Chunk as RagChunk

    doc_id = _uuid_mod.uuid4().hex
    doc_count = 0
    for idx, chunk in enumerate(chunks):
        if not chunk.content.strip():
            continue
        embedding: list[float] = []
        if embedder:
            try:
                embeddings = await embed_texts([chunk.content], provider=embedder)
                embedding = embeddings[0]
            except Exception:
                pass
        rag_chunk = RagChunk(
            document_id=doc_id,
            content=chunk.content,
            embedding=embedding,
            chunk_index=idx,
            metadata={**{k: str(v) for k, v in metadata.items()}, "source_type": body.source_type},
        )
        try:
            store.ingest_chunk(rag_chunk, collection_id=body.collection_id, tenant_ctx=tenant_ctx)
            doc_count += 1
        except Exception:
            pass

    return {
        "collection_id": body.collection_id,
        "source_url": body.url,
        "source_type": body.source_type,
        "chunks_ingested": doc_count,
        "total_chars": len(content),
    }


# ---------------------------------------------------------------------------
# Endpoints — structured source ingestors (PDF, DOCX, GitHub, Confluence, etc.)
# ---------------------------------------------------------------------------

async def _ingest_chunks_from_source(
    store: KnowledgeStore,
    chunks: list[dict],
    collection_id: str,
    tenant_ctx: Any,
    embedder: Any,
) -> int:
    """Embed and ingest a list of chunk dicts returned by an ingestor."""
    from app.providers.base import embed_texts
    ingested = 0
    for chunk_data in chunks:
        content = chunk_data.get("content", "")
        if not content.strip():
            continue
        embedding: list[float] = []
        if embedder:
            try:
                embeddings = await embed_texts([content], provider=embedder)
                embedding = embeddings[0]
            except Exception:
                pass
        rag_chunk = Chunk(
            document_id=_uuid.uuid4().hex,
            content=content,
            embedding=embedding,
            chunk_index=ingested,
            metadata={
                k: str(v) for k, v in (chunk_data.get("metadata") or {}).items()
            } | {
                "source_url": chunk_data.get("source_url", ""),
                "source_type": chunk_data.get("source_type", ""),
                "source_doc_id": chunk_data.get("source_doc_id", ""),
                "page_number": str(chunk_data.get("page_number") or ""),
            },
        )
        try:
            store.ingest_chunk(rag_chunk, collection_id=collection_id, tenant_ctx=tenant_ctx)
            ingested += 1
        except Exception:
            pass
    return ingested


@router.post("/ingest/pdf", status_code=201)
async def ingest_pdf(
    request: Request,
    file: UploadFile = File(...),
    collection_id: str = Form(...),
    source_url: str = Form(default=""),
) -> dict[str, Any]:
    """Ingest a PDF file into a knowledge collection with page-level citation metadata."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    content_bytes = await file.read()
    filename = file.filename or "uploaded.pdf"

    from app.knowledge.ingestors.pdf_ingestor import PdfIngestor
    ingestor = PdfIngestor()
    chunks = ingestor.extract_chunks(
        content=content_bytes,
        filename=filename,
        source_url=source_url or f"file://{filename}",
    )

    ingested = await _ingest_chunks_from_source(store, chunks, collection_id, tenant, embedder)
    return {"chunks_ingested": ingested, "source": filename, "source_type": "pdf"}


@router.post("/ingest/docx", status_code=201)
async def ingest_docx(
    request: Request,
    file: UploadFile = File(...),
    collection_id: str = Form(...),
    source_url: str = Form(default=""),
) -> dict[str, Any]:
    """Ingest a DOCX file into a knowledge collection."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    content_bytes = await file.read()
    filename = file.filename or "uploaded.docx"

    from app.knowledge.ingestors.docx_ingestor import DocxIngestor
    ingestor = DocxIngestor()
    chunks = ingestor.extract_chunks(
        content=content_bytes,
        filename=filename,
        source_url=source_url or f"file://{filename}",
    )

    ingested = await _ingest_chunks_from_source(store, chunks, collection_id, tenant, embedder)
    return {"chunks_ingested": ingested, "source": filename, "source_type": "docx"}


@router.post("/ingest/github", status_code=202)
async def ingest_github(
    request: Request, body: GitHubIngestRequest
) -> dict[str, Any]:
    """Ingest a GitHub repository into a knowledge collection via GitHub REST API."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    from app.knowledge.ingestors.github_ingestor import GitHubIngestor
    ingestor = GitHubIngestor()
    chunks = await ingestor.ingest_repo(
        body.owner, body.repo,
        branch=body.branch,
        max_files=body.max_files,
    )

    ingested = await _ingest_chunks_from_source(store, chunks, body.collection_id, tenant, embedder)
    return {
        "chunks_ingested": ingested,
        "source": f"github:{body.owner}/{body.repo}",
        "source_type": "github",
    }


@router.post("/ingest/confluence", status_code=202)
async def ingest_confluence(
    request: Request, body: ConfluenceIngestRequest
) -> dict[str, Any]:
    """Ingest a Confluence space into a knowledge collection."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    from app.knowledge.ingestors.confluence_ingestor import ConfluenceIngestor
    ingestor = ConfluenceIngestor(
        base_url=body.base_url,
        token=body.token.get_secret_value(),  # SecretStr: extract only at point of use
        user=body.user,
    )
    chunks = await ingestor.ingest_space(body.space_key, max_pages=body.max_pages)

    ingested = await _ingest_chunks_from_source(store, chunks, body.collection_id, tenant, embedder)
    return {
        "chunks_ingested": ingested,
        "source": f"confluence:{body.space_key}",
        "source_type": "confluence",
    }


@router.post("/ingest/jira", status_code=202)
async def ingest_jira(
    request: Request, body: JiraIngestRequest
) -> dict[str, Any]:
    """Ingest Jira project issues into a knowledge collection."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    from app.knowledge.ingestors.jira_ingestor import JiraIngestor
    ingestor = JiraIngestor(
        base_url=body.base_url,
        token=body.token.get_secret_value(),  # SecretStr: extract only at point of use
        user=body.user,
    )
    chunks = await ingestor.ingest_project(
        body.project_key,
        jql_extra=body.jql_extra,
        max_issues=body.max_issues,
    )

    ingested = await _ingest_chunks_from_source(store, chunks, body.collection_id, tenant, embedder)
    return {
        "chunks_ingested": ingested,
        "source": f"jira:{body.project_key}",
        "source_type": "jira",
    }


@router.post("/ingest/slack", status_code=202)
async def ingest_slack(
    request: Request, body: SlackIngestRequest
) -> dict[str, Any]:
    """Ingest a Slack channel's message history into a knowledge collection."""
    tenant = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)

    from app.knowledge.ingestors.slack_ingestor import SlackIngestor
    ingestor = SlackIngestor(token=body.token.get_secret_value())
    chunks = await ingestor.ingest_channel(
        body.channel_id,
        channel_name=body.channel_name,
        max_messages=body.max_messages,
    )

    ingested = await _ingest_chunks_from_source(store, chunks, body.collection_id, tenant, embedder)
    return {
        "chunks_ingested": ingested,
        "source": f"slack:{body.channel_id}",
        "source_type": "slack",
    }


# ---------------------------------------------------------------------------
# H-7: Federated search across multiple collections
# ---------------------------------------------------------------------------

@router.post("/search/federated")
async def federated_search_endpoint(
    request: Request,
    body: dict[str, Any],
) -> dict[str, Any]:
    """Search across multiple knowledge collections with score normalization."""
    tenant_ctx: TenantContext = _require_tenant(request)
    store = _knowledge_store(request)
    embedder = getattr(request.app.state, "embedder", None)
    if embedder is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="No embedding provider configured",
        )

    query: str = body.get("query", "")
    collection_ids: list[str] = body.get("collection_ids", [])
    top_k: int = int(body.get("top_k", 10))
    top_k = max(1, min(100, top_k))

    if not query:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="query is required")
    if not collection_ids:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="collection_ids is required")

    from app.knowledge.federated_search import federated_search
    results = await federated_search(
        query=query,
        collection_ids=collection_ids,
        store=store,
        top_k=top_k,
    )
    return {
        "results": results,
        "total": len(results),
        "collections_searched": len(collection_ids),
    }
