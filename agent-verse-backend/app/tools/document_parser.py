"""Document parsing tools for PDF, CSV, DOCX, and plain text files.

All parsing is done locally — no cloud services required.
Dependencies (optional — graceful degradation if missing):
  pip install pypdf python-docx chardet
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from app.observability.logging import get_logger

logger = get_logger(__name__)

_MAX_PAGES = 50
_MAX_CHARS = 100_000


@dataclass
class ParsedDocument:
    filename: str
    content: str
    page_count: int | None
    metadata: dict[str, Any]
    format: str
    truncated: bool = False


class DocumentParserTool:
    """Parse documents from various formats into text.

    Supports: PDF, CSV, DOCX, TXT, MD, JSON, YAML
    All processing is local — no cloud APIs.
    """

    name = "parse_document"
    description = "Parse documents (PDF, CSV, DOCX, TXT) into readable text."

    async def execute(
        self,
        *,
        file_path: str = "",
        content_bytes: bytes | None = None,
        filename: str = "document",
    ) -> dict:
        import asyncio
        return await asyncio.get_running_loop().run_in_executor(
            None, self._parse_sync, file_path, content_bytes, filename
        )

    def _parse_sync(
        self, file_path: str, content_bytes: bytes | None, filename: str
    ) -> dict:
        # Determine format
        name = filename or (Path(file_path).name if file_path else "doc")
        ext = Path(name).suffix.lower()

        if content_bytes is None and file_path:
            try:
                with open(file_path, "rb") as f:
                    content_bytes = f.read()
            except Exception as exc:
                return {"error": f"Cannot read file: {exc}"}

        if content_bytes is None:
            return {"error": "No content provided"}

        parsers = {
            ".pdf": self._parse_pdf,
            ".csv": self._parse_csv,
            ".docx": self._parse_docx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
            ".json": self._parse_json,
            ".yaml": self._parse_yaml,
            ".yml": self._parse_yaml,
        }
        parser = parsers.get(ext, self._parse_text)

        try:
            doc = parser(content_bytes, filename=name)
            return {
                "filename": doc.filename,
                "format": doc.format,
                "page_count": doc.page_count,
                "content": doc.content,
                "metadata": doc.metadata,
                "truncated": doc.truncated,
                "char_count": len(doc.content),
            }
        except Exception as exc:
            logger.warning("document_parse_failed", filename=name, error=str(exc))
            return {"error": str(exc), "filename": name}

    def _parse_pdf(self, data: bytes, *, filename: str) -> ParsedDocument:
        try:
            import pypdf  # type: ignore[import]
        except ImportError:
            try:
                from PyPDF2 import PdfReader  # type: ignore[import]
                reader = PdfReader(io.BytesIO(data))
                pages = []
                for i, page in enumerate(reader.pages[:_MAX_PAGES]):
                    pages.append(page.extract_text() or "")
                text = "\n\n".join(pages)
                trunc = len(reader.pages) > _MAX_PAGES
                return ParsedDocument(
                    filename=filename,
                    content=text[:_MAX_CHARS],
                    page_count=len(reader.pages),
                    metadata={},
                    format="pdf",
                    truncated=trunc or len(text) > _MAX_CHARS,
                )
            except ImportError:
                pass
            return ParsedDocument(
                filename=filename,
                content="[PDF parsing unavailable: install pypdf]",
                page_count=None,
                metadata={},
                format="pdf",
            )

        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = []
        for i, page in enumerate(reader.pages[:_MAX_PAGES]):
            pages.append(page.extract_text() or "")
        text = "\n\n".join(pages)
        trunc = len(reader.pages) > _MAX_PAGES
        return ParsedDocument(
            filename=filename,
            content=text[:_MAX_CHARS],
            page_count=len(reader.pages),
            metadata=dict(reader.metadata or {}),
            format="pdf",
            truncated=trunc or len(text) > _MAX_CHARS,
        )

    def _parse_csv(self, data: bytes, *, filename: str) -> ParsedDocument:
        import csv
        try:
            import chardet
            enc = chardet.detect(data)["encoding"] or "utf-8"
        except ImportError:
            enc = "utf-8"

        text = data.decode(enc, errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        # Format as markdown table
        if not rows:
            return ParsedDocument(
                filename=filename,
                content="[Empty CSV]",
                page_count=None,
                metadata={"rows": 0},
                format="csv",
            )

        header = rows[0]
        divider = "|" + "|".join(["---"] * len(header)) + "|"
        lines = ["|" + "|".join(header) + "|", divider]
        for row in rows[1:1001]:  # max 1000 rows
            lines.append("|" + "|".join(str(c) for c in row) + "|")

        content = "\n".join(lines)
        return ParsedDocument(
            filename=filename,
            content=content[:_MAX_CHARS],
            page_count=None,
            metadata={"rows": len(rows) - 1, "columns": len(header), "headers": header},
            format="csv",
            truncated=len(rows) > 1001 or len(content) > _MAX_CHARS,
        )

    def _parse_docx(self, data: bytes, *, filename: str) -> ParsedDocument:
        try:
            from docx import Document  # type: ignore[import]
        except ImportError:
            return ParsedDocument(
                filename=filename,
                content="[DOCX parsing unavailable: install python-docx]",
                page_count=None,
                metadata={},
                format="docx",
            )

        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return ParsedDocument(
            filename=filename,
            content=content[:_MAX_CHARS],
            page_count=None,
            metadata={"paragraph_count": len(paragraphs)},
            format="docx",
            truncated=len(content) > _MAX_CHARS,
        )

    def _parse_text(self, data: bytes, *, filename: str) -> ParsedDocument:
        text = data.decode("utf-8", errors="replace")
        return ParsedDocument(
            filename=filename,
            content=text[:_MAX_CHARS],
            page_count=None,
            metadata={},
            format="text",
            truncated=len(text) > _MAX_CHARS,
        )

    def _parse_json(self, data: bytes, *, filename: str) -> ParsedDocument:
        import json
        text = data.decode("utf-8", errors="replace")
        try:
            obj = json.loads(text)
            formatted = json.dumps(obj, indent=2)
        except Exception:
            formatted = text
        return ParsedDocument(
            filename=filename,
            content=formatted[:_MAX_CHARS],
            page_count=None,
            metadata={},
            format="json",
            truncated=len(formatted) > _MAX_CHARS,
        )

    def _parse_yaml(self, data: bytes, *, filename: str) -> ParsedDocument:
        try:
            import yaml  # type: ignore[import]
            text = data.decode("utf-8", errors="replace")
            obj = yaml.safe_load(text)
            formatted = yaml.dump(obj, default_flow_style=False)
        except Exception:
            formatted = data.decode("utf-8", errors="replace")
        return ParsedDocument(
            filename=filename,
            content=formatted[:_MAX_CHARS],
            page_count=None,
            metadata={},
            format="yaml",
            truncated=len(formatted) > _MAX_CHARS,
        )

    def to_tool_def(self) -> dict:
        return {
            "name": self.name,
            "description": self.description,
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Absolute path to the file"},
                    "filename": {"type": "string", "description": "Filename with extension (for format detection)"},
                },
            },
        }
